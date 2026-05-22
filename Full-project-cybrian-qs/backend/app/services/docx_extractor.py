"""
Structured DOCX extraction preserving headings, lists, tables, and document order.
"""
from __future__ import annotations

import re
from io import BytesIO
from typing import Any, Dict, List, Tuple

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from .pdf_extractor import _clean_line, sanitize_extracted_text

def _heading_level_from_style(style_name: str) -> int | None:
    name = (style_name or "").strip().lower()
    if not name:
        return None
    if name in {"title", "document title", "titel"}:
        return 1
    m = re.search(r"heading\s*(\d+)", name)
    if m:
        return max(1, min(6, int(m.group(1))))
    if name.startswith("heading") and name[-1:].isdigit():
        return max(1, min(6, int(name[-1])))
    if name in {"berschrift", "überschrift 1", "überschrift 2", "überschrift 3"}:
        digits = re.search(r"(\d+)", name)
        return max(1, min(6, int(digits.group(1)))) if digits else 2
    return None


def _paragraph_is_list_item(paragraph: Paragraph) -> bool:
    p_pr = paragraph._p.pPr  # noqa: SLF001
    if p_pr is None:
        return False
    return p_pr.numPr is not None


def _list_style(paragraph: Paragraph) -> str:
    """Return 'bullet' or 'numbered' for a Word list paragraph."""
    text = paragraph.text.strip()
    if re.match(r"^[-*•]\s+", text):
        return "bullet"
    if re.match(r"^\d+[\.)]\s+", text):
        return "numbered"
    return "numbered"


def _strip_list_prefix(text: str) -> str:
    text = _clean_line(text)
    text = re.sub(r"^[-*•]\s+", "", text)
    text = re.sub(r"^\d+[\.)]\s+", "", text)
    return text


def _table_to_block(table: Table) -> Dict[str, Any] | None:
    rows: List[List[str]] = []
    for row in table.rows:
        cells = [_clean_line(cell.text) for cell in row.cells]
        if any(cells):
            rows.append(cells)
    if not rows:
        return None
    return {"type": "table", "rows": rows}


def _flush_list(buffer: List[str], list_kind: str | None, blocks: List[Dict[str, Any]]) -> None:
    if not buffer or not list_kind:
        return
    items = [_strip_list_prefix(item) for item in buffer if _strip_list_prefix(item)]
    if not items:
        buffer.clear()
        return
    block_type = "numbered_list" if list_kind == "numbered" else "bullet_list"
    blocks.append({"type": block_type, "items": items})
    buffer.clear()


def _paragraph_to_blocks(paragraph: Paragraph) -> List[Dict[str, Any]]:
    text = _clean_line(paragraph.text)
    if not text:
        return []

    level = _heading_level_from_style(paragraph.style.name if paragraph.style else "")
    if level is not None:
        block_type = "section_heading" if level <= 2 else "heading"
        return [{"type": block_type, "text": text, "level": min(3, level)}]

    if _paragraph_is_list_item(paragraph):
        return [{"type": "_list_item", "text": text, "list_style": _list_style(paragraph)}]

    return [{"type": "paragraph", "text": text}]


def _iter_body_elements(document: Document):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def extract_docx_bytes(docx_bytes: bytes) -> Tuple[List[Dict[str, Any]], str]:
    """
    Extract DOCX into typed blocks (headings, lists, tables, paragraphs) in document order.
    """
    document = Document(BytesIO(docx_bytes))
    blocks: List[Dict[str, Any]] = []
    list_buffer: List[str] = []
    list_kind: str | None = None

    for element in _iter_body_elements(document):
        if isinstance(element, Table):
            _flush_list(list_buffer, list_kind, blocks)
            list_kind = None
            table_block = _table_to_block(element)
            if table_block:
                blocks.append(table_block)
            continue

        for piece in _paragraph_to_blocks(element):
            if piece.get("type") == "_list_item":
                style = piece.get("list_style") or "bullet"
                item_text = piece.get("text") or ""
                if list_kind and list_kind != style:
                    _flush_list(list_buffer, list_kind, blocks)
                list_kind = style
                list_buffer.append(item_text)
                continue

            _flush_list(list_buffer, list_kind, blocks)
            list_kind = None
            blocks.append({k: v for k, v in piece.items() if not k.startswith("_")})

    _flush_list(list_buffer, list_kind, blocks)

    from .document_structure import refine_blocks

    text = _blocks_to_plain_text(blocks)
    blocks = refine_blocks(blocks, text)
    return blocks, sanitize_extracted_text(text)


def _blocks_to_plain_text(blocks: List[Dict[str, Any]]) -> str:
    parts: List[str] = []
    for block in blocks or []:
        btype = str(block.get("type", "")).lower()
        if btype in {"section_heading", "heading", "paragraph"}:
            val = str(block.get("text", "")).strip()
            if val:
                parts.append(val)
        elif btype in {"bullet_list", "numbered_list"}:
            for item in block.get("items") or []:
                if str(item).strip():
                    parts.append(str(item).strip())
        elif btype == "table":
            for row in block.get("rows") or []:
                cells = [str(c).strip() for c in row or [] if str(c).strip()]
                if cells:
                    parts.append(" | ".join(cells))
        elif btype == "two_column_row":
            left = str(block.get("left", "")).strip()
            right = str(block.get("right", "")).strip()
            if left and right:
                parts.append(f"{left}: {right}")
    return "\n\n".join(parts).strip()


# ──────────────────────────────────────────────────────────────────────────────
# Sequential Elements Extraction — unified format for instant frontend loading
# ──────────────────────────────────────────────────────────────────────────────

def extract_docx_elements(docx_bytes: bytes) -> List[Dict[str, Any]]:
    """
    Extract DOCX in perfect reading order using pre-mapped XML element references.

    Pre-builds two O(1) lookup dicts before the body walk:
      p_map   : {paragraph._p  → Paragraph}  — avoids repeated python-docx wrapping
      tbl_map : {table._tbl    → Table}

    This sidesteps the O(N²) performance cliff caused by python-docx re-wrapping
    XML nodes on every access when iterating large documents sequentially.

    Returns unified elements list:
      {"type": "text",  "style": "heading"|"paragraph", "content": "<str>"}
      {"type": "table", "content": [["H1","H2"], ["R1C1","R1C2"]]}
    """
    document = Document(BytesIO(docx_bytes))

    # Pre-map native XML → python-docx wrappers for O(1) per child
    p_map: Dict[Any, Paragraph] = {p._p: p for p in document.paragraphs}  # noqa: SLF001
    tbl_map: Dict[Any, Table] = {t._tbl: t for t in document.tables}  # noqa: SLF001

    elements: List[Dict[str, Any]] = []

    for child in document.element.body.iterchildren():
        tag = child.tag

        if tag == qn("w:p"):
            # Resolve via pre-map (O(1)); fall back to fresh wrap if not found
            para: Paragraph = p_map.get(child) or Paragraph(child, document)
            text = _clean_line(para.text)
            if not text:
                continue
            level = _heading_level_from_style(para.style.name if para.style else "")
            if level is not None:
                elements.append({"type": "text", "style": "heading", "content": text})
            else:
                elements.append({"type": "text", "style": "paragraph", "content": text})

        elif tag == qn("w:tbl"):
            # Resolve via pre-map (O(1)); fall back to fresh wrap if not found
            tbl: Table = tbl_map.get(child) or Table(child, document)
            rows: List[List[str]] = []
            for row in tbl.rows:
                cells = [_clean_line(cell.text) for cell in row.cells]
                if any(cells):
                    rows.append(cells)
            if rows:
                elements.append({"type": "table", "content": rows})

    return elements
