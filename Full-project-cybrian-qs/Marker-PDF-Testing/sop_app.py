import streamlit as st
import pdfplumber
import fitz  # PyMuPDF
from docx import Document
import pandas as pd
import io

st.set_page_config(layout="wide", page_title="Universal Sequential SOP Renderer")

# Custom Document UI CSS Styling to mimic actual document structures
st.markdown("""
    <style>
    .sop-canvas {
        background-color: #ffffff;
        color: #222222;
        padding: 40px;
        border-radius: 6px;
        box-shadow: 0px 4px 15px rgba(0,0,0,0.08);
        border: 1px solid #e0e0e0;
        margin-bottom: 25px;
    }
    .sop-text-block {
        font-family: 'Courier New', Courier, monospace;
        font-size: 13px;
        line-height: 1.6;
        white-space: pre-wrap;
        margin-bottom: 8px;
    }
    .page-divider {
        color: #1565c0;
        font-weight: bold;
        border-bottom: 2px dashed #1565c0;
        padding-bottom: 3px;
        margin-top: 30px;
        margin-bottom: 15px;
        font-family: sans-serif;
    }
    </style>
""", unsafe_allow_html=True)

st.title("📄 Universal Sequential SOP Layout Engine")
st.caption("Chronologically interleaves headings, paragraphs, and tables for both PDF and DOCX documents.")

uploaded_file = st.file_uploader("Upload SOP File (.pdf or .docx)", type=["pdf", "docx"])

# --- ADVANCED FIXED SEQUENTIAL DOCX PARSER ---
def extract_docx_in_reading_order(file_bytes):
    """
    Parses a DOCX file by traversing its raw XML body tree layout chronologically.
    Maps underlying OpenXML block nodes safely to correct text and matrix table components.
    """
    doc = Document(io.BytesIO(file_bytes))
    sequential_elements = []
    
    # Pre-map the native element tree references to bypass python-docx lookup limitations
    paragraph_map = {p._p: p for p in doc.paragraphs}
    table_map = {t._tbl: t for t in doc.tables}
    
    # Walk through the parent body XML container element by element
    for child in doc.element.body:
        # Match element if it's an OpenXML Paragraph entry (w:p)
        if child in paragraph_map:
            p_obj = paragraph_map[child]
            if p_obj.text.strip():
                sequential_elements.append({
                    "type": "text",
                    "content": p_obj.text
                })
        
        # Match element if it's an OpenXML Table entry (w:tbl)
        elif child in table_map:
            t_obj = table_map[child]
            table_matrix = []
            for row in t_obj.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                table_matrix.append(row_cells)
            
            if table_matrix:
                sequential_elements.append({
                    "type": "table",
                    "content": table_matrix
                })
                
    return sequential_elements


# --- HIGH-FIDELITY PDF PARSER ---
def extract_pdf_in_reading_order(file_bytes):
    pages_data = []
    
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for p_idx, page in enumerate(pdf.pages):
            constructed_lines = []
            
            # 1. Isolate tables and track bounding coordinates
            tables_found = page.find_tables(table_settings={
                "vertical_strategy": "lines",
                "horizontal_strategy": "lines",
                "snap_tolerance": 4
            })
            
            table_zones = []
            for t in tables_found:
                table_zones.append({
                    "top": t.bbox[1],
                    "bottom": t.bbox[3],
                    "cells": t.extract()
                })
            
            # 2. Extract words and group into rows
            text_objects = page.extract_words()
            lines_map = {}
            for word in text_objects:
                y_coord = round(word["top"], 1)
                found_match = False
                for existing_y in lines_map.keys():
                    if abs(existing_y - y_coord) < 4:
                        lines_map[existing_y].append(word)
                        found_match = True
                        break
                if not found_match:
                    lines_map[y_coord] = [word]
            
            for y_coord, words in lines_map.items():
                sorted_words = sorted(words, key=lambda w: w["x0"])
                line_text = " ".join([w["text"] for w in sorted_words])
                
                is_inside_table = False
                for zone in table_zones:
                    if zone["top"] <= y_coord <= zone["bottom"]:
                        is_inside_table = True
                        break
                
                if not is_inside_table:
                    constructed_lines.append({
                        "type": "text",
                        "top": y_coord,
                        "content": line_text
                    })
            
            # 3. Add table zones
            for zone in table_zones:
                constructed_lines.append({
                    "type": "table",
                    "top": zone["top"],
                    "content": zone["cells"]
                })
            
            # 4. Sort everything vertically
            chronological_elements = sorted(constructed_lines, key=lambda e: e["top"])
            pages_data.append({"page_num": p_idx + 1, "elements": chronological_elements})
            
    return pages_data


# --- STREAMLIT PIPELINE USER INTERFACE ---
if uploaded_file:
    file_bytes = uploaded_file.read()
    file_type = uploaded_file.name.split(".")[-1].lower()
    
    if file_type == "pdf":
        if st.button("Render PDF Layout Flow"):
            with st.spinner("Analyzing PDF structural vectors..."):
                structured_data = extract_pdf_in_reading_order(file_bytes)
                
                for page in structured_data:
                    st.markdown(f"<div class='page-divider'>📄 ORIGINAL PDF PAGE {page['page_num']}</div>", unsafe_allow_html=True)
                    st.markdown("<div class='sop-canvas'>", unsafe_allow_html=True)
                    
                    for idx, element in enumerate(page["elements"]):
                        if element["type"] == "text":
                            st.markdown(f"<div class='sop-text-block'>{element['content']}</div>", unsafe_allow_html=True)
                        elif element["type"] == "table":
                            raw_table = element["content"]
                            filtered_table = [row for row in raw_table if any(c is not None and str(c).strip() != "" for c in row)]
                            if filtered_table:
                                df = pd.DataFrame(filtered_table)
                                df.columns = [f"Col {i}" if v is None or str(v).strip() == "" else str(v).replace('\n', ' ') for i, v in enumerate(df.iloc[0])]
                                df = df.drop(df.index[0]).reset_index(drop=True)
                                st.dataframe(df, width="stretch", key=f"pdf_tbl_{page['page_num']}_{idx}")
                    st.markdown("</div>", unsafe_allow_html=True)
                    
    elif file_type == "docx":
        if st.button("Render DOCX Layout Flow"):
            with st.spinner("Analyzing Word document XML nodes..."):
                structured_data = extract_docx_in_reading_order(file_bytes)
                
                st.markdown(f"<div class='page-divider'>📝 ORIGINAL WORD DOCUMENT SEQUENCE</div>", unsafe_allow_html=True)
                st.markdown("<div class='sop-canvas'>", unsafe_allow_html=True)
                
                for idx, element in enumerate(structured_data):
                    if element["type"] == "text":
                        st.markdown(f"<div class='sop-text-block'>{element['content']}</div>", unsafe_allow_html=True)
                    elif element["type"] == "table":
                        raw_table = element["content"]
                        filtered_table = [row for row in raw_table if any(c is not None and str(c).strip() != "" for c in row)]
                        if filtered_table:
                            df = pd.DataFrame(filtered_table)
                            df.columns = [f"Col {i}" if v is None or str(v).strip() == "" else str(v).replace('\n', ' ') for i, v in enumerate(df.iloc[0])]
                            df = df.drop(df.index[0]).reset_index(drop=True)
                            st.dataframe(df, width="stretch", key=f"docx_tbl_{idx}")
                st.markdown("</div>", unsafe_allow_html=True)